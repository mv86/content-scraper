"""Convert web pages into docx documents for key-word analysis."""
from docx import Document
from docx.shared import Inches

from document.create_template import create_template


def create_document(filename, url, stop_words):
    """Create document template from provided arguments and convert to docx document. 

       Save to path: /tmp/FILENAME
    """
    doc = Document()
    template = create_template(url, stop_words)

    doc.add_heading(template.heading, 0)

    doc.add_paragraph().add_run('Title:').bold = True
    doc.add_paragraph(template.title)

    doc.add_paragraph().add_run('Description:').bold = True
    doc.add_paragraph(template.description)

    doc.add_paragraph().add_run('Top Ten Keywords:').bold = True
    _add_top_ten_words_table(doc, template.top_ten_words)
    # Add blank paragraph under table
    doc.add_paragraph()

    doc.add_paragraph().add_run('Content:').bold = True
    _add_page_content(doc, template.content)

    path = f'/tmp/{filename}'
    doc.save(path)


def _add_page_content(doc, page_content):
    """Helper function for create_document."""
    for item in page_content:
        # Unpack tuple
        tag, content = item
        # If header tag add text and header type
        if tag[0] == 'h':
            header = f'{content} - ({tag})'
            doc.add_heading(header, level=1)
        # If any list item add bullet point
        elif tag == 'li':
            doc.add_paragraph(content, style='List Bullet')
        # Else paragraph, no formatting needed
        else:
            doc.add_paragraph(content)


def _add_top_ten_words_table(doc, top_ten_words):
    """Helper function for create_document."""
    table = doc.add_table(rows=1, cols=3, style='Table Grid')
    _set_column_widths(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = 'Keyword'
    hdr_cells[2].text = 'Frequency'

    for idx, item in enumerate(top_ten_words, start=1):
        word, count = item
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = word
        row_cells[2].text = str(count)


# Need to edit column width and every cell inside for width to stick. See below:
# https://stackoverflow.com/questions/43051462/python-docx-how-to-set-cell-width-in-tables
def _set_column_widths(table):
    """Helper function for _add_top_ten_words_table."""
    widths = (Inches(0.8), Inches(1.6), Inches(1))

    for column, width in zip(table.columns, widths):
        column.width = width

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
